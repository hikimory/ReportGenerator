import io
import yaml
import numpy as np
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt

def load_yaml(filename):
    with open(filename, 'r', encoding='utf-8') as file:
        return yaml.safe_load(file)

def plot_histogram(subject_name, periods, legend):
    # Data preparation
    high_pct = []
    sufficient_pct = []
    intermediate_pct = []
    low_pct = []
    not_certified_pct = []
    period_names = []

    # Calculation
    for period in periods:
        period_names.append(period['name'])
        total = period['total']
        high_pct.append((period['highLevel'] / total) * 100)
        sufficient_pct.append((period['sufficientLevel'] / total) * 100)
        intermediate_pct.append((period['intermediateLevel'] / total) * 100)
        low_pct.append((period['lowLevel'] / total) * 100)
        not_certified_pct.append((period['notCertified'] / total) * 100)
    
    # Creating a diagram
    bar_width = 0.15
    x = np.arange(len(period_names))  # Места для каждого периода
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Adding columns for each level
    ax.bar(x - 2*bar_width, high_pct, bar_width, label="Высокий уровень", color=legend['highLevel'])
    ax.bar(x - bar_width, sufficient_pct, bar_width, label="Достаточный уровень", color=legend['sufficientLevel'])
    ax.bar(x, intermediate_pct, bar_width, label="Средний уровень", color=legend['intermediateLevel'])
    ax.bar(x + bar_width, low_pct, bar_width, label="Низкий уровень", color=legend['lowLevel'])
    ax.bar(x + 2*bar_width, not_certified_pct, bar_width, label="Не аттестованный", color=legend['notCertified'])
    
    # Chart settings
    # ax.set_title(f'{subject_name} - Percentage of Students by Level')
    # ax.set_xlabel('Period')
    ax.set_ylabel("Процентное соотношение")
    all_numbers = high_pct + sufficient_pct + intermediate_pct + low_pct + not_certified_pct 
    max_number = max(all_numbers)
    plt.yticks(np.arange(0, max_number + 1, 5))
    plt.grid(True, axis='y')

    ax.set_xticks(x)
    ax.set_xticklabels(period_names)
    ax.legend()

    plt.tight_layout()
    histogram_filename = f'histogram_{subject_name}.png'
    plt.savefig(histogram_filename)
    plt.close()
    return histogram_filename

def insert_histogram_into_docx(doc, subject_name, histogram_filename, width, height):
    doc.add_paragraph(f"{subject_name}")
    doc.add_paragraph()
    doc.add_picture(histogram_filename, width=Inches(width), height=Inches(height))

def main():
    width = 14.8 / 2.54
    height = 8.9 / 2.54
    doc = Document()
    yaml_filename = 'config.yaml'
    data = load_yaml(yaml_filename)

    for subject in data['subjects']:
        hist_filename = plot_histogram(subject['name'], subject['periods'], data['legend'])
        insert_histogram_into_docx(doc, subject['name'], hist_filename, width, height)

    doc.save('subject_charts.docx')
    print("Документ успешно создан!")

if __name__ == '__main__':
    main()
